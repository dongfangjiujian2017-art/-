import os
import io
import streamlit as st
import requests
import json
import time

# 尝试导入必要库
try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    st.error("❌ 缺少依赖库：python-docx。请运行 `pip install python-docx`。")
    st.stop()

# --- 页面配置 ---
st.set_page_config(page_title="DeepSeek x Word 智能助手", page_icon="🐼", layout="wide")

# --- 持久化配置管理 ---
CONFIG_FILE = "config.json"

def load_config():
    if "saved_api_key" in st.session_state:
        return {"api_key": st.session_state.saved_api_key, "base_url": st.session_state.get("saved_base_url", "https://api.deepseek.com")}
    env_key = os.environ.get("DEEPSEEK_API_KEY", "")
    if env_key:
        return {"api_key": env_key, "base_url": "https://api.deepseek.com"}
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r") as f:
                return json.load(f)
        except:
            pass
    return {"api_key": "", "base_url": "https://api.deepseek.com"}

def save_config(api_key, base_url):
    st.session_state.saved_api_key = api_key
    st.session_state.saved_base_url = base_url
    try:
        with open(CONFIG_FILE, "w") as f:
            json.dump({"api_key": api_key, "base_url": base_url}, f)
    except:
        pass 

# --- 核心逻辑：深度 Word 替换 ---

def _replace_in_paragraph(paragraph, replace_dict):
    """替换段落中的文本并保留格式"""
    for old_text, new_text in replace_dict.items():
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

def _process_shapes(doc_part, replace_dict):
    """
    递归处理文档部分中的所有形状（包括文本框）。
    注意：此函数尝试捕获存在于页眉/页脚/正文中的浮动文本框。
    """
    # 扫描所有形状
    for shape in doc_part.inline_shapes:
        pass # inline_shapes 通常不含文本框

    # 获取底层 XML 元素中的文本框
    # 这里通过访问 doc_part 的 element 查找所有的 <w:t>
    # 但为了更稳妥地保留格式，我们处理文本框内部的段落
    try:
        # 获取该部分所有的文本内容块（包含文本框内部）
        for p in doc_part.paragraphs:
            _replace_in_paragraph(p, replace_dict)
            
        # 遍历所有表格
        for table in doc_part.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph(p, replace_dict)
    except:
        pass

def deep_smart_replace(doc, replace_dict):
    """
    深度扫描：正文、表格、页眉、页脚、文本框。
    """
    # 1. 替换正文中的所有内容
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replace_dict)
    
    # 2. 替换正文表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replace_dict)

    # 3. 替换页眉和页脚（包含其中的文本框）
    for section in doc.sections:
        # 处理页眉
        header = section.header
        for p in header.paragraphs:
            _replace_in_paragraph(p, replace_dict)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph(p, replace_dict)
        
        # 处理页脚
        footer = section.footer
        for p in footer.paragraphs:
            _replace_in_paragraph(p, replace_dict)
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph(p, replace_dict)

    # 4. 关键：扫描隐藏在文本框（Shapes）里的内容
    # 由于 python-docx 对 Shape 的 API 支持有限，
    # 我们通过直接遍历文档的 body element 寻找所有的 txbxContent (文本框容器)
    for txbx in doc.element.xpath('//w:txbxContent'):
        # 在文本框内部寻找段落
        for p_element in txbx.xpath('.//w:p'):
            from docx.text.paragraph import Paragraph
            p = Paragraph(p_element, doc)
            _replace_in_paragraph(p, replace_dict)
            
    return doc

def parse_replace_text(text):
    replace_dict = {}
    for line in text.split('\n'):
        clean_line = line.replace('**', '').replace('`', '').strip()
        if "==>>" in clean_line:
            parts = clean_line.split("==>>")
            if len(parts) == 2:
                replace_dict[parts[0].strip()] = parts[1].strip()
    return replace_dict

# --- AI 逻辑 ---
def get_deepseek_rules(api_key, base_url, doc_content, user_demand):
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    prompt = f"""
    你是一个文档处理专家。下面是一段 Word 模板的内容：
    ---
    {doc_content}
    ---
    用户的修改需求是："{user_demand}"
    请对比模板内容，找出需要被替换的精确原文字（包括页码占位符），并生成替换列表。
    格式要求：旧内容 ==>> 新内容。禁止 Markdown 格式。
    """
    data = {"model": "deepseek-chat", "messages": [{"role": "system", "content": "你是一个专业的文档分析助手。"}, {"role": "user", "content": prompt}], "stream": False}
    try:
        response = requests.post(f"{base_url}/chat/completions", headers=headers, json=data, timeout=60)
        return response.json()['choices'][0]['message']['content']
    except:
        return "❌ API 连接超时。"

# --- UI 界面 ---
def main():
    st.title("🐼 DeepSeek x Word 智能助手 (深度替换版)")
    config = load_config()
    
    with st.sidebar:
        st.header("⚙️ 设置")
        input_key = st.text_input("DeepSeek API Key", value=config["api_key"], type="password")
        input_url = st.text_input("Base URL", value=config["base_url"])
        if st.button("💾 临时记住 Key"):
            save_config(input_key, input_url)
            st.success("已记住设置")
        st.divider()
        st.info("💡 此版本已支持替换【文本框】及【页眉页脚】内的内容。")

    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("1. 上传模板")
        uploaded_file = st.file_uploader("选择 Word 文件", type=["docx"])
        doc_sample = ""
        if uploaded_file:
            doc = Document(uploaded_file)
            doc_sample = "\n".join([p.text for p in doc.paragraphs])[:3000]
            st.success(f"✅ 已加载: {uploaded_file.name}")

        st.subheader("2. 描述需求")
        st.write("如果你想去掉页脚文本框里的页码，请写：'把 第 PAGE 页 共 NUMPAGES 页 替换为空'")
        user_demand = st.text_area("告诉 AI 你想改什么？")
        
        if "ai_rules" not in st.session_state:
            st.session_state.ai_rules = ""

        if st.button("✨ 生成替换规则", type="primary"):
            if not input_key: st.error("请先输入 API Key")
            elif not uploaded_file: st.warning("请上传文件")
            else:
                with st.spinner("DeepSeek 分析中..."):
                    result = get_deepseek_rules(input_key, input_url, doc_sample, user_demand)
                    st.session_state.ai_rules = result

    with col2:
        st.subheader("3. 执行替换")
        rules_text = st.text_area("最终替换规则", value=st.session_state.ai_rules, height=300)
        
        if st.button("🚀 深度扫描并下载", use_container_width=True):
            if not uploaded_file: st.error("请上传模板")
            else:
                replacements = parse_replace_text(rules_text)
                with st.spinner("正在递归搜索文本框、页眉、页脚..."):
                    doc = Document(uploaded_file)
                    # 调用深度替换函数
                    processed_doc = deep_smart_replace(doc, replacements)
                    
                    output = io.BytesIO()
                    processed_doc.save(output)
                    output.seek(0)
                    st.download_button("📥 下载处理后的 Word", data=output, file_name=f"Deep_Fixed_{uploaded_file.name}", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

if __name__ == "__main__":
    main()
